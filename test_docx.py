import logging
import random
import re
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.shared import Inches, Cm, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.run import Run
from docx.opc.constants import RELATIONSHIP_TYPE


logging.basicConfig(filename='logs.log', level=logging.INFO)

def set_cell_value(cell, text=None, image_path=None, heading=None, bold=False, isHyperlink=False, hyperlink_url=None):
    if not re.search(r'<wp:inline', cell._element.xml):
        if text is not None and image_path is None:
            if not isHyperlink:
                if heading is not None:
                    cell_text = cell.paragraphs[0].add_run(heading)
                    cell_text.font.name = 'Calibri'
                    cell_text.font.size = Pt(16)
                    cell_text.bold = bold
                    cell_text.add_break()
                cell_text = cell.paragraphs[0].add_run(text)
                cell_text.font.name = 'Calibri'
                cell_text.font.size = Pt(12)
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                cell_text.add_break()
                logging.info(f'Set cell value to "{text}"')
            else:
                if heading is not None:
                    cell_text = cell.paragraphs[0].add_run(heading)
                    cell_text.font.name = 'Calibri'
                    cell_text.font.size = Pt(16)
                    cell_text.bold = bold
                    cell_text.add_break()
                cell_text = cell.paragraphs[0].add_run(text)
                cell_text.font.name = 'Calibri'
                cell_text.font.size = Pt(12)
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # Add hyperlink styling
                r_id = cell.part.relate_to(hyperlink_url, RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
                hyperlink = OxmlElement('w:hyperlink')
                hyperlink.set(qn('r:id'), r_id)
                hyperlink.append(cell_text._element)
                hyperlinkRPr = OxmlElement('w:rPr')
                hyperlinkRPr.set(qn('w:u'), 'single')
                hyperlinkRPr.set(qn('w:color'), '0072c6')
                hyperlink.append(hyperlinkRPr)
                cell.paragraphs[0]._p.append(hyperlink)
                cell_text.add_break()
                
                logging.info(f'Set cell value to hyperlink "{text}"')
        elif text is None and image_path is not None:
            cell.width = Inches(3)
            cell.height = Inches(1.5)
            paragraph = cell.paragraphs[0]
            run = paragraph.add_run()
            run.add_picture(image_path, width=Inches(2.5))
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run.add_break()
            logging.info(f'Set cell value to image "{image_path}"')
    else:
        logging.error('Cell has inline element, cannot set value')





def create_header(doc, header_text):
    header = doc.sections[0].header
    paragraph = header.paragraphs[0]
    paragraph.text = header_text
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


def create_footer(doc, footer_text):
    footer = doc.sections[0].footer
    paragraph = footer.paragraphs[0]
    paragraph.text = footer_text
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.distance_from_top = Cm(1.0)


def create_table(doc, rows=3, cols=2, table_name="table_name", cell_values=None):
    logging.info(f'Creating table with {rows} rows and {cols} columns')
    table = doc.add_table(rows=rows+1, cols=cols)
    table.autofit = False
    table.allow_autofit = False
    table.style = 'Table Grid'
    widths = [Inches(3) for _ in range(cols)]

    # set table header row
    header_row = table.rows[0]
    for cell, width in zip(header_row.cells, widths):
        cell.width = width
    header_row.height = Inches(0.3)
    header_row.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    header_row.cells[-1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # merge cells in header row
    merge_cell = header_row.cells[0].merge(header_row.cells[-1])
    set_cell_value(merge_cell, table_name, bold=True)
    merge_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    for row_idx, row in enumerate(table.rows[1:]):
        row.height = Inches(0.3)
        for col_idx, width in enumerate(widths):
            cell = row.cells[col_idx]
            cell.width = width
            if cell_values and row_idx < len(cell_values) and col_idx < len(cell_values[row_idx]):
                cell_value = cell_values[row_idx][col_idx]
                if isinstance(cell_value, tuple):
                    heading, value, img_path, add_image, is_hyperlink, hyperlink_url = cell_value
                    if heading:
                        value = f"{value}"
                else:
                    value = cell_value
                set_cell_value(cell, value, heading=heading, isHyperlink=is_hyperlink, hyperlink_url=hyperlink_url)



                
                if add_image and img_path:
                    cell.width = widths[0]
                    cell.height = Inches(1.5)
                    paragraph = cell.paragraphs[0]
                    run = paragraph.add_run()
                    run.add_picture(img_path, width=Inches(1.5))
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER




def create_heading(doc):
    logging.info('Creating heading')
    doc.add_heading('Heading', level=1)
    add_new_paragraph(doc)


def create_bullet_point_list(doc):
    logging.info('Creating bullet point list')
    num_items = random.randint(3, 5)
    for i in range(num_items):
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f'Bullet point {i + 1}')
    add_new_paragraph(doc)


def create_numbered_list(doc):
    logging.info('Creating numbered list')
    num_items = random.randint(3, 5)
    for i in range(num_items):
        p = doc.add_paragraph(style='List Number')
        p.add_run(f'Numbered item {i + 1}')
    add_new_paragraph(doc)


def create_nested_bullet_point_list(doc):
    logging.info('Creating nested bullet point list')
    num_items = random.randint(2, 4)
    for i in range(num_items):
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f'Bullet point {i + 1}')
        if random.random() < 0.5:
            nested_num_items = random.randint(2, 4)
            for j in range(nested_num_items):
                p = doc.add_paragraph(style='List Bullet 2')
                p.add_run(f'Nested bullet point {i + 1}.{j + 1}')
    add_new_paragraph(doc)


def create_nested_numbered_list(doc):
    logging.info('Creating nested numbered list')
    num_items = random.randint(2, 4)
    for i in range(num_items):
        p = doc.add_paragraph(style='List Number')
        p.add_run(f'Numbered item {i + 1}')
        if random.random() < 0.5:
            nested_num_items = random.randint(2, 4)
            for j in range(nested_num_items):
                p = doc.add_paragraph(style='List Number 2')
                p.add_run(f'Nested numbered item {i + 1}.{j + 1}')
    add_new_paragraph(doc)

def section(doc):
    doc.add_paragraph('---')


def add_new_paragraph(doc, text=None, add_link=False, link_text=None, link_url=None):
    """
    A function that adds a new paragraph to a document.

    Args:
        doc (docx.document.Document): The document to add the paragraph to.
        text (str): The text content of the new paragraph.
        add_link (bool): A flag indicating whether to add a hyperlink to the paragraph.
        link_text (str): The text content of the hyperlink (if add_link is True).
        link_url (str): The URL of the hyperlink (if add_link is True).
    """

    if add_link and (link_text is None or link_url is None):
        raise ValueError("link_text and link_url cannot be None if add_link is True")

    if not add_link:
        paragraph = doc.add_paragraph()
        paragraph.add_run(text)
    else:
        paragraph = doc.add_paragraph(text)
        add_hyperlink(paragraph, link_text, link_url)



def add_hyperlink(paragraph, text, url):
    """
    A function that places a hyperlink within a paragraph object.

    Args:
        paragraph (docx.text.paragraph.Paragraph): The paragraph to add the hyperlink to.
        text (str): The display text of the hyperlink.
        url (str): The URL to link to.
    
    Returns:
        A reference to the newly created hyperlink object.
    """

    part = paragraph.part
    r_id = part.relate_to(url, RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    # Create the w:hyperlink tag and add needed values
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id, )

    # Create a w:r element
    new_run = OxmlElement('w:r')

    # Create a new w:rPr element
    rPr = OxmlElement('w:rPr')

    # Add hyperlink styling
    rStyle = OxmlElement('w:rStyle')
    rStyle.set(qn('w:val'), 'Hyperlink')
    rPr.append(rStyle)
    
    # Add underlining
    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)

    # Set font color to blue
    color = OxmlElement('w:color')
    color.set(qn('w:val'), '0000FF')
    rPr.append(color)

    # Join all the xml elements together add add the required text to the w:r element
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)

    paragraph._p.append(hyperlink)

    return hyperlink



def insert_image(doc, image_path='assets/image.png'):
    doc.add_picture(image_path, width=Inches(10), height=Inches(3))


cell_values = [
    [(None, "Hello" , "assets/image.png", True, False, ""), ("Response", "There", "", False, True, "https://www.example.com")],
    [("Mood", "General", "", False, False, ""), ("Character", "Kenobi", "assets/image2.png", True, True, "https://www.example.com")],
    [("Sound", "Roaring", "assets/image2.png", True, False, ""), ("Pet", "Kitten", "", False, True, "https://www.example.com")]
]




arg_funcs = {
    'aa': lambda doc: create_header(doc, 'Header'),
    'bb': lambda doc: create_footer(doc, 'Footer'),
    'cc': lambda doc: create_table(doc, rows=len(cell_values), cols=max([len(row) for row in cell_values]), cell_values=cell_values, table_name="Card"),
    'dd': create_heading,
    'ee': lambda doc: create_table(doc, rows=3, cols=2, add_image=True, table_name='table_name'),
    'ff': create_bullet_point_list,
    'gg': create_numbered_list,
    'hh': create_nested_bullet_point_list,
    'ii': create_nested_numbered_list,
    'jj': lambda doc: insert_image(doc, image_path='assets/image.png'),
    'kk': lambda doc: section(doc),
    'mm': lambda doc: add_new_paragraph(doc, add_link=True, link_text='Click here', link_url='https://www.example.com'),
    'nn': lambda doc: add_new_paragraph(doc, '', add_link=True, link_text='Click here', link_url='https://www.example.com'),
    'oo': lambda doc: create_table(doc, rows=len(cell_values), cols=max([len(row) for row in cell_values]), cell_values=cell_values, table_name="Card")
}


def create_doc(args):
    logging.info(f'Creating document with arguments: {args}')
    doc = Document()

    for arg in args:
        func = arg_funcs.get(arg)
        if func is not None:
            func(doc)
        else:
            logging.warning(f'Invalid argument "{arg}"')
    return doc


if __name__ == '__main__':
    import sys

    args = sys.argv[1:]
    print(f"Input arguments: {args}")
    doc = create_doc(args)
    doc.save('document.docx')
    print("Document created successfully")